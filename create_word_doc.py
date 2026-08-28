import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_element(name):
    return OxmlElement(name)

def set_cell_border(cell, **kwargs):
    """
    kwargs: top, bottom, left, right
    values: dict(sz=12, val='single', color='FF0000', space='0')
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = create_element('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        if border_name in kwargs:
            edge = create_element(f'w:{border_name}')
            for key, val in kwargs[border_name].items():
                edge.set(qn(f'w:{key}'), str(val))
            tcBorders.append(edge)
    tcPr.append(tcBorders)

def add_header(doc):
    p_title1 = doc.add_paragraph()
    p_title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title1.paragraph_format.space_before = Pt(0)
    p_title1.paragraph_format.space_after = Pt(2)
    run1 = p_title1.add_run("মুন্সীগঞ্জ জেলা ছাত্র সংগঠন")
    run1.font.name = 'Kalpurush'
    run1.font.size = Pt(22)
    run1.font.bold = True
    run1.font.color.rgb = RGBColor(11, 44, 92) # Deep navy blue

    p_title2 = doc.add_paragraph()
    p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title2.paragraph_format.space_before = Pt(0)
    p_title2.paragraph_format.space_after = Pt(2)
    run2 = p_title2.add_run("Munshiganj District Student’s Association")
    run2.font.name = 'Calibri'
    run2.font.size = Pt(16)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(11, 44, 92)

    p_title3 = doc.add_paragraph()
    p_title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title3.paragraph_format.space_before = Pt(0)
    p_title3.paragraph_format.space_after = Pt(6)
    run3 = p_title3.add_run("গোপালগঞ্জ বিজ্ঞান ও প্রযুক্তি বিশ্ববিদ্যালয়")
    run3.font.name = 'Kalpurush'
    run3.font.size = Pt(16)
    run3.font.bold = True
    run3.font.color.rgb = RGBColor(11, 44, 92)

    # Double bottom border line simulation or horizontal line
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_before = Pt(0)
    p_line.paragraph_format.space_after = Pt(12)
    pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                     r'<w:bottom w:val="double" w:sz="12" w:space="1" w:color="0B2C5C"/>'
                     r'</w:pBdr>')
    p_line._p.get_or_add_pPr().append(pBdr)

    # Date line on the right
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.paragraph_format.space_before = Pt(0)
    p_date.paragraph_format.space_after = Pt(18)
    run_date = p_date.add_run("তারিখঃ ")
    run_date.font.name = 'Kalpurush'
    run_date.font.size = Pt(14)
    run_date.font.bold = True

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    run.font.name = 'Kalpurush'
    run.font.size = Pt(20)
    run.font.bold = True

def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'Kalpurush'
    run.font.size = Pt(15)
    run.font.bold = True

def add_body_p(doc, text, bold_prefix="", space_after=8, space_before=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Kalpurush'
        r_bold.font.size = Pt(13)
        r_bold.font.bold = True
        
    r_text = p.add_run(text)
    r_text.font.name = 'Kalpurush'
    r_text.font.size = Pt(13)

doc = Document()

# Set Margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Normal Style default font
style = doc.styles['Normal']
font = style.font
font.name = 'Kalpurush'
font.size = Pt(13)

# --- PAGE 1 ---
add_header(doc)
add_heading_1(doc, "গঠনতন্ত্র")
add_section_title(doc, "ভূমিকা (Introduction):")
add_body_p(doc, "মুন্সীগঞ্জ জেলা থেকে গোপালগঞ্জ বিজ্ঞান ও প্রযুক্তি বিশ্ববিদ্যালয়ে ভর্তি হওয়া শিক্ষার্থীদের পারস্পরিক সহযোগিতা ও ভ্রাতৃত্বপূর্ণ সম্পর্ক বিনির্মাণের লক্ষ্যে সম্পূর্ণ অরাজনৈতিক, অলাভজনক, অসাম্প্রদায়িক ও সেচ্ছাসেবী সামাজিক এই সংগঠনটি প্রতিষ্ঠিত হলো।", space_after=14)

add_section_title(doc, "ধারা-১: নাম (Name)")
add_body_p(doc, " এই সংগঠনের নাম হবে \"মুন্সীগঞ্জ জেলা ছাত্র সংগঠন,, গোপালগঞ্জ বিজ্ঞান ও প্রযুক্তি বিশ্ববিদ্যালয়\" সংক্ষেপে \"মুন্সীগঞ্জ জেলা ছাত্র সংগঠন, গোবিপ্রবি।\"", bold_prefix="(ক)")
add_body_p(doc, "অথবা,", space_before=2, space_after=2)
add_body_p(doc, " ইংরেজিতে \"Munshiganj District Students' Association, Gopalganj Science and Technology University \" সংক্ষেপে \"Munshiganj District Students' Association, GSTU\"", bold_prefix="(খ)", space_after=14)

add_section_title(doc, "ধারা-২: প্রতীকচিহ্ন (Logo)")
add_body_p(doc, " এই সংগঠনের একটি নির্ধারিত প্রতীকচিহ্ন (লোগো) থাকবে। লোগোর মধ্যে মুন্সীগঞ্জ জেলার ম্যাপ, উন্মুক্ত বই ও শিখা, শাপলা, পাতার তোড়া বৃত্ত দ্বারা আবদ্ধ লোগোটির বৃত্তের মধ্যে সংগঠনের পূর্ণনাম এবং বিশ্ববিদ্যালয়ের সংক্ষিপ্ত নাম ইংরেজিতে লিখিত থাকবে।", bold_prefix="(ক)")
add_body_p(doc, " সংগঠনের প্রতীকচিহ্ন (লোগো) নির্বাহী সদস্যদের অনুমোদন সাপেক্ষে পরিবর্তন যোগ্য। এক্ষেত্রে অত্র গঠনতন্ত্রের ধারা-২ এর (ক) সংশোধন বাধ্যতামূলক।", bold_prefix="(খ)", space_after=14)

# --- PAGE 2 ---
doc.add_page_break()
add_header(doc)
add_section_title(doc, "ধারা-৩: উদ্দেশ্য (Objectives)")
add_body_p(doc, "এই সংগঠনের উদ্দেশ্যসমূহ নিম্নরূপ-", space_after=8)
add_body_p(doc, " মুন্সীগঞ্জ জেলার উচ্চ মাধ্যমিক সার্টিফিকেট (এইচএসসি) পরীক্ষায় উত্তীর্ণ শিক্ষার্থীদের দেশের বিভিন্ন বিশ্ববিদ্যালয়ের ভর্তি পরীক্ষার জন্য প্রয়োজনীয় দিকনির্দেশনা, পরামর্শ ও সহায়তা প্রদান।", bold_prefix="(ক)")
add_body_p(doc, " গোপালগঞ্জ বিজ্ঞান ও প্রযুক্তি বিশ্ববিদ্যালয়ে ভর্তিচ্ছু শিক্ষার্থীদের নির্ভরযোগ্য তথ্য সরবরাহ, একাডেমিক গাইডলাইন প্রদান এবং প্রেষণামূলক কার্যক্রমের মাধ্যমে সহযোগিতা প্রদান।", bold_prefix="(খ)")
add_body_p(doc, " গোপালগঞ্জ বিজ্ঞান ও প্রযুক্তি বিশ্ববিদ্যালয়ে নবভর্তি শিক্ষার্থীদের সাথে অত্র বিশ্ববিদ্যালয়ে অধ্যয়নরত শিক্ষার্থীদের মধ্যে পরিচিতি স্থাপন এবং পারস্পরিক সৌহার্দ্যপূর্ণ সম্পর্ক গড়ে তোলা।", bold_prefix="(গ)")
add_body_p(doc, " অত্র বিশ্ববিদ্যালয়ে অধ্যয়নরত শিক্ষার্থীদের মধ্যে ভ্রাতৃত্ববোধ, পারস্পরিক সহযোগিতা ও সামাজিক দায়িত্ববোধ জাগ্রত করা এবং তা সুদৃঢ়করণ।", bold_prefix="(ঘ)")
add_body_p(doc, " গোপালগঞ্জ বিজ্ঞান ও প্রযুক্তি বিশ্ববিদ্যালয় থেকে শিক্ষা সমাপ্ত করা মুন্সীগঞ্জ জেলার সাবেক শিক্ষার্থীদের (অ্যালামনাই) সাথে অধ্যয়নরত শিক্ষার্থীদের মধ্যে যোগাযোগ ও সমন্বয় প্রতিষ্ঠার মাধ্যমে একটি কার্যকর নেটওয়ার্ক গড়ে তোলা।", bold_prefix="(ঙ)")
add_body_p(doc, " সংগঠনের মূলনীতির উপর ভিত্তি করে গৃহীত ঘোষণাপত্র ও কর্মসূচী বাস্তবায়ন করা।", bold_prefix="(চ)", space_after=14)

# --- PAGE 3 ---
doc.add_page_break()
add_header(doc)
add_section_title(doc, "ধারা-৪: সদস্য (Member)")
add_body_p(doc, " ধর্ম, বর্ণ, গোত্র, দলমত নির্বিশেষে বাংলাদেশের মহান মুক্তিযুদ্ধ ও সার্বভৌমত্বের প্রতি শ্রদ্ধাশীল সংগঠনের গঠনতন্ত্রের প্রতি বিশ্বাসী বিক্রমপুর তথা মুন্সিগঞ্জ জেলার অধিবাসী (জন্ম বা পৈতৃকনিবাস সূত্রে) এবং গোপালগঞ্জ বিজ্ঞান ও প্রযুক্তি বিশ্ববিদ্যালয়ে অধ্যয়নরত যে কোনো শিক্ষার্থী সমিতির সদস্য হতে পারবেন।", bold_prefix="ক)", space_after=14)

add_section_title(doc, "ধারা-৫: উপদেষ্টা পরিষদ (Advisory Council)")
add_body_p(doc, " উপদেষ্টা পরিষদ শিক্ষক উপদেষ্টা এবং বিশেষ উপদেষ্টাদের সমন্বয়ে গঠিত হবে।", bold_prefix="(ক)")
add_body_p(doc, " গোপালগঞ্জ বিজ্ঞান ও প্রযুক্তি বিশ্ববিদ্যালয়ের যেসকল শিক্ষক জন্মসূত্রে কিংবা পৈতৃক নিবাস সূত্রে মুন্সীগঞ্জের নিবাসী শুধুমাত্র তারাই এই সমিতির শিক্ষক উপদেষ্টা হিসেবে বিবেচিত হবেন।", bold_prefix="(খ)")
add_body_p(doc, " সংগঠনের সকল সাবেক সভাপতি, সহ-সভাপতি, সাধারণ সম্পাদক এবং সমিতির প্রতিষ্ঠাকালীন সময়ে গোপালগঞ্জ বিজ্ঞান ও প্রযুক্তি বিশ্ববিদ্যালয়ের মুন্সীগঞ্জ নিবাসী যে সকল সাবেক শিক্ষার্থী সংগঠনকে সার্বিকভাবে সহায়তা করবেন কেবল তারাই এই সমিতির বিশেষ উপদেষ্টা হিসেবে বিবেচিত হবেন।", bold_prefix="(গ)")
add_body_p(doc, " যে সকল সদস্য সংগঠনের যেকোনো কার্যনির্বাহী কমিটির একাধিক দায়িত্ব পালন করবে তারাও বিশেষ উপদেষ্টা বলে বিবেচিত হবেন।", bold_prefix="(ঘ)")
add_body_p(doc, "উপদেষ্টা পরিষদ হবে সর্বোচ্চ দিক নির্দেশক পরিষদ। কার্যনির্বাহী কমিটি যথাযথ ভাবে কার্যক্রম পরিচালনায় ব্যর্থ হলে উপদেষ্টা কমিটি হস্তক্ষেপ করবেন এবং সঠিক দিক নির্দেশনা দিবেন।", bold_prefix="(ঙ)", space_after=14)

# --- PAGE 4 ---
doc.add_page_break()
add_header(doc)
add_section_title(doc, "ধারা-৬: কার্যনির্বাহী কমিটি ও গঠন প্রক্রিয়া (Executive Committee and Formation Process):")
add_body_p(doc, " প্রত্যেক কার্যনির্বাহী কমিটির মেয়াদ হবে ১ বছর।", bold_prefix="(ক)")
add_body_p(doc, " কার্যনির্বাহী কমিটির সভাপতি, সহ-সভাপতি, সাধারণ সম্পাদক যেহেতু তাদের মতামতের ভিত্তিতে মনোনীত হবেন কিংবা নির্বাচিত হবেন সেহেতু পরবর্তীতে কোন অবস্থাতেই তাদের পক্ষ থেকে আপত্তি পত্র গ্রহণ করা হবেনা।", bold_prefix="(খ)")
add_body_p(doc, " নির্বাচনের মাধ্যমে কার্যনির্বাহী কমিটি গঠন করার প্রয়োজন হলে বিদ্যমান সদস্যদের মধ্য থেকে উপদেষ্টাদের মতামতের ভিত্তিতে এবং সুষ্ঠু ও নিরপেক্ষ নির্বাচনের মাধ্যমে গণতান্ত্রিক উপায়ে কার্যনির্বাহী কমিটি গঠন করা হবে।", bold_prefix="(গ)")
add_body_p(doc, " পদ সমপরিমাণ সদস্য থাকা অবস্থায় বিশেষ উপদেষ্টাদের নেতৃত্বে এবং বিদ্যমান সদস্যদের আলোচনার মাধ্যমে জ্যেষ্ঠতার ভিত্তিতে পদবণ্টনের মাধ্যমে কার্যনির্বাহী কমিটি গঠন করা হবে।", bold_prefix="(ঘ)")
add_body_p(doc, " সংগঠনের কার্যনির্বাহী কমিটি সংগঠনের প্রয়োজনে কোন ব্যক্তিকে কার্যনির্বাহী কমিটির একাধিক (সর্বোচ্চ ২টি) দপ্তরের দায়িত্ব প্রদান করা হয় তাহলে তিনি দায়িত্ব গ্রহণ করবেন।", bold_prefix="(ঙ)")
add_body_p(doc, " নির্বাচনের মাধ্যমে কমিটি গঠন করার সময় সংগঠনের সকল সদস্য নির্বাচনে অংশ নিতে পারবেন।", bold_prefix="(চ)", space_after=14)

# --- PAGE 5 ---
doc.add_page_break()
add_header(doc)
add_section_title(doc, "ধারা-৭: কার্যনির্বাহী কমিটির দপ্তর সমূহ (Executive Committee Offices)")
add_body_p(doc, "সংগঠনের সভাপতি সংগঠনের প্রধান প্রতিনিধির মর্যাদা পাবেন। সভাপতি সংগঠনের সভায় সভাপতিত্ব করবেন। কার্যনির্বাহী কমিটি কর্তৃক গৃহীত যাবতীয় সিদ্ধান্ত অনুমোদনের দায়িত্ব সভাপতির উপর থাকবে। সভাপতি সভা আহ্বানের জন্য সাধারণ সম্পাদককে নির্দেশ দিবেন। সাধারণ সম্পাদক নির্দেশ অমান্য করলে তিনি নিজে সভা আহ্বান করতে পারবেন।", bold_prefix="(ক) সভাপতি:\n")
add_body_p(doc, "সহ-সভাপতির সকল কার্যক্রমে সভাপতির সাহায্যকারীর ভূমিকা পালন করবেন। সভাপতির অনুপস্থিতিতে বা অন্য কোন কারনে সভাপতি স্বাভাবিক কার্যক্রম পরিচালনায় অপারগ হইলে সহ-সভাপতি, সভাপতির দায়িত্ব পালন করবেন।", bold_prefix="(খ) সহ-সভাপতি\n")
add_body_p(doc, "সভাপতির নির্দেশে অথবা পরামর্শে সাধারণ সম্পাদক সভা আহ্বান করবেন। অন্যান্য সম্পাদক ও সদস্যদের সমন্বয় করার মাধ্যমে সংগঠনের কার্যক্রম এগিয়ে নিয়ে যাবেন। সমমনা সংগঠনগুলোর সাথে যোগাযোগ রক্ষা করবেন।", bold_prefix="(গ) সাধারণ সম্পাদক\n")
add_body_p(doc, "তিনি সংগঠনকে সাংগঠনিকভাবে সুদৃঢ় করার জন্য সবসময় সচেষ্ট থাকবেন। কার্যনির্বাহী কমিটি কর্তৃক গৃহীত যে কোন সাংগঠনিক সিদ্ধান্ত বাস্তবায়নে তিনি অগ্রণী ভূমিকা পালন করবেন।", bold_prefix="(ঘ) সাংগঠনিক সম্পাদক:\n", space_after=14)

# --- PAGE 6 ---
doc.add_page_break()
add_header(doc)
add_body_p(doc, "", bold_prefix="(ঙ)কোষাধ্যক্ষ:", space_after=2)
add_body_p(doc, " সংগঠনের আয় বায় এর হিসাব রাখবেন এবং তহবিল সংগ্রহের জন্য সচেষ্ট থাকবেন।", bold_prefix="(১)")
add_body_p(doc, " অর্থের উৎস হিসেবে সদস্য চাঁদা, উপদেষ্টা মণ্ডলীর চাঁদা এবং প্রাক্তনদের চাঁদা সংগ্রহ করবেন।", bold_prefix="(২)")
add_body_p(doc, " কার্যনির্বাহী কমিটির নির্ধারিত সময়কালের মধ্যে অর্থ রিপোর্ট (আয়-ব্যয় বিবরণী) উত্থাপন করবেন।", bold_prefix="(৩)", space_after=10)

add_body_p(doc, "কার্যকরী কমিটির সকল প্রকার জিনিসপত্র দেখাশোনা ও সংরক্ষণের দায়িত্ব তাহার উপর থাকবে।\nসংগঠনের সকল প্রেস বিজ্ঞপ্তি তার স্বাক্ষরে প্রকাশিত হবে। একইসাথে তিনি সেসকল প্রেস বিজ্ঞপ্তির ব্যাপারে দায়বদ্ধ থাকবেন।", bold_prefix="(চ) দপ্তর সম্পাদক:\n", space_after=10)

add_body_p(doc, "সংগঠনের প্রয়োজনীয় তথ্য সংগ্রহ ও সংরক্ষণের দায়িত্ব তার উপর থাকবে। সমিতির প্রচার কার্যক্রম তাহার দ্বারা সম্পন্ন হইবে। সমিতির যাবতীয় প্রকাশনার দায়িত্ব তাহার উপর থাকবে।", bold_prefix="(ছ) তথ্য, প্রচার ও প্রকাশনা সম্পাদক:\n", space_after=10)

add_body_p(doc, "কার্যনির্বাহী কমিটি কর্তৃক গৃহীত শিক্ষামূলক যে কোন উদ্যোগ বাস্তবায়নের মুখ্য দায়িত্ব তাহার উপর থাকবে। যে কোন প্রকার সাহিত্য প্রতিযোগিতা বা সাহিত্যমূলক আলোচনা আয়োজন করার মুখ্য দায়িত্ব তাহার উপর থাকবে। সংগঠন কর্তৃক যাবতীয় গবেষণা কাজের মুখ্য দায়িত্ব তাহার উপর থাকবে।", bold_prefix="(জ) শিক্ষা ও গবেষণা সম্পাদক:\n", space_after=10)

add_body_p(doc, "সংগঠনের সদস্যদের মধ্যে তথ্যপ্রযুক্তি ও বিজ্ঞান বিষয়ে সচেতনতা বৃদ্ধি ও প্রায়োগিক দিক নিশ্চিত করার লক্ষ্যে কাজ করবেন।", bold_prefix="(ঝ) বিজ্ঞান ও প্রযুক্তি বিষয়ক সম্পাদক:\n", space_after=14)

# --- PAGE 7 ---
doc.add_page_break()
add_header(doc)
add_body_p(doc, "সংগঠনের আয়োজিত যাবতীয় ক্রীড়া প্রতিযোগিতা তাহার তত্ত্বাবধানে অনুষ্ঠিত হইবে। বন্ধুত্বপূর্ণ ক্রীড়া প্রতিযোগিতার জন্য অন্য যে কোন সংগঠনের সহিত যোগাযোগ করা এবং সাফল্যমণ্ডিত করা তাহার দায়িত্ব। যাবতীয় ক্রীড়া ও সাংস্কৃতিক সরঞ্জামাদি সংগ্রহ ও সংরক্ষণ করা তার দায়িত্ব।", bold_prefix="(ঞ) ক্রীড়া সম্পাদক:\n", space_after=10)

add_body_p(doc, "সংগঠন কর্তৃক আয়োজিত বিভিন্ন অনুষ্ঠান ও কর্মসূচী বাস্তবায়নে তিনি সচেষ্ট থাকবেন। অনুষ্ঠান ও কর্মসূচী সম্পন্ন হবার পরবর্তী সাধারণ সভায় উক্ত আয়োজন সম্পর্কে রিপোর্ট প্রদান করবেন।", bold_prefix="(ট) অনুষ্ঠান ও কর্মসূচী বিষয়ক সম্পাদকঃ\n", space_after=10)

add_body_p(doc, "সংগঠনের সকল প্রকার সাংস্কৃতিক অনুষ্ঠান উদযাপন করা, আয়োজন করা ও পরিচালনা করা তাহার দায়িত্ব।", bold_prefix="(ঠ) সাংস্কৃতিক সম্পাদক:\n", space_after=10)

add_body_p(doc, "সামাজিক কল্যাণমূলক যাবতীয় কার্যাবলী সম্পাদনের দায়িত্ব তাহার উপর ন্যস্ত থাকবে। সংগঠনের সদস্যদের মধ্যে সামাজিক মূল্যবোধ সৃষ্টি ও পরিবেশ বিষয়ে সচেতনতা বৃদ্ধির প্রয়াস।", bold_prefix="(ড) সমাজকল্যাণ ও পরিবেশ বিষয়ক সম্পাদক:\n", space_after=14)

# --- PAGE 8 ---
doc.add_page_break()
add_header(doc)
add_section_title(doc, "ধারা-৮: সভা ও কার্যক্রম (Meetings and Activities)")
add_body_p(doc, " প্রতি দুই মাস অন্তর একবার কার্যনির্বাহী কমিটির সভা অনুষ্ঠিত হবে।", bold_prefix="১.")
add_body_p(doc, " সভা অনুষ্ঠিত হবার ৭ দিন পূর্বে সাধারণ সম্পাদক কার্যনির্বাহী কমিটির সকল সদস্যকে অবহিত করবেন।", bold_prefix="২.")
add_body_p(doc, " ২৪ ঘণ্টার বিজ্ঞপ্তিতে জরুরী সভা আহ্বান করা যাবে।", bold_prefix="৩.")
add_body_p(doc, " সভার বিজ্ঞপ্তি অবশ্যই সংগঠনের ফেসবুক পেইজ এ দিতে হবে।", bold_prefix="৪.")
add_body_p(doc, " যে কোন সভায় মোট সদস্যের এক তৃতীয়াংশ সদস্য উপস্থিতি তা কোরাম বলে গণ্য হবে। এর কম হলে সভা কার্যকর হবে না।", bold_prefix="৫.")
add_body_p(doc, " সভায় সিদ্ধান্ত গ্রহণের ক্ষেত্রে সংখ্যাগরিষ্ঠতাই হবে মাপকাঠি।", bold_prefix="৬.")
add_body_p(doc, " এক চতুর্থাংশ সদস্য সভাপতির নিকট লিখিত ভাবে দাবি জানালে জরুরী সভা আহ্বান করা যাবে।", bold_prefix="৭.")
add_body_p(doc, " পূর্ববর্তী সভার গৃহীত সিদ্ধান্ত এবং তৎপরবর্তী কর্মকাণ্ডের বিষয়ে সাংগঠনিক সম্পাদক রিপোর্ট প্রকাশ করবেন", bold_prefix="৮.")
add_body_p(doc, " প্রতিটি সভায় সমিতির পরবর্তী কর্মসূচী ও পরিকল্পনা আলোচিত হবে।", bold_prefix="৯.")
add_body_p(doc, " সভায় সাংগঠনিক সম্পাদক সাংগঠনিক রিপোর্ট এবং কোষাধ্যক্ষ অবশ্যই অর্থ রিপোর্ট প্রদান করবেন।", bold_prefix="১০.", space_after=14)

# --- PAGE 9 ---
doc.add_page_break()
add_header(doc)
add_section_title(doc, "ধারা-৯: অর্থ ব্যবস্থাপনা ও তহবিল (Financial Management and Fund)")
add_body_p(doc, " সংগঠনের তহবিল গঠিত হবে সদস্যদের অবদান, অনুদান, উৎস বা অন্যান্য সেচ্ছাসেবী উৎস থেকে।", bold_prefix="ক.")
add_body_p(doc, " সংগঠনের উপদেষ্টামণ্ডলীর সদস্যরা সমিতির তহবিলে বার্ষিক অন্যুন পাঁচশত (৫০০) টাকা অবদান রাখবেন।", bold_prefix="খ.")
add_body_p(doc, " তহবিল ব্যবস্থাপনার জন্য একজন সম্পাদক থাকবেন এবং সাধারণ সম্পাদকের অনুমোদনক্রমে তহবিল ব্যবহার করবেন।", bold_prefix="গ.")
add_body_p(doc, " সংগঠনের অর্থ ব্যয় হবে শুধুমাত্র গঠনতন্ত্রে নির্ধারিত উদ্দেশ্য এবং কার্যক্রমের জন্য।", bold_prefix="ঘ.")
add_body_p(doc, " সমস্ত আর্থিক লেনদেন স্বচ্ছ, নিয়মিত এবং নথিভুক্ত হতে হবে।", bold_prefix="ঙ.", space_after=14)

add_section_title(doc, "ধারা-১০: শৃঙ্খলা ও নিয়ম ভঙ্গ (Discipline and Misconduct)")
add_body_p(doc, " সংগঠনের সকল স্তরের সদস্যকে গঠনতন্ত্র, কমিটির নিয়ম ও নীতিমালা যথাযথভাবে মেনে চলতে হবে।", bold_prefix="ক.")
add_body_p(doc, " সদস্যরা যে কোনও প্রকার শৃঙ্খলা ভঙ্গ, অসদাচরণ বা গঠনতন্ত্র লঙ্ঘন করলে কার্যনির্বাহী কমিটি প্রয়োজনীয় ব্যবস্থা গ্রহণ করবে।", bold_prefix="খ.")
add_body_p(doc, " কোনো সদস্য দেশের ফৌজদারি অপরাধে দোষী প্রমাণিত হলে তার সদস্যপদ স্থায়ীভাবে বাতিল হবে। এক্ষেত্রে সদস্যপদ পুনর্বহাল যোগ্য নয়।", bold_prefix="গ.")
add_body_p(doc, " কোনো সদস্য বিশ্ববিদ্যালয় থেকে অস্থায়ী কিংবা স্থায়ী বহিষ্কারের শাস্তি পেলে, সেই সদস্যকে সংগঠন থেকে স্থায়ীভাবে বহিষ্কার করা হবে। এরূপ বহিষ্কার আদেশ প্রত্যাহারযোগ্য নয়।", bold_prefix="ঘ.", space_after=14)

# --- PAGE 10 ---
doc.add_page_break()
add_header(doc)
add_body_p(doc, " ফৌজদারি অপরাধে দোষী সাব্যস্ত হওয়ার ফলে সদস্যপদ বাতিল হওয়া কিংবা বিশ্ববিদ্যালয়ের বহিষ্কার আদেশের কারণে সংগঠন থেকে বহিষ্কার হওয়া সদস্যদের বাহিরে অন্য কোনো সদস্যকে শৃঙ্খলা ভঙ্গের কারণে সংগঠন থেকে স্থায়ী বহিষ্কার করতে হলে সকল উপদেষ্টা, নির্বাহী সদস্য এবং সম্পাদকীয় সদস্যের অনুমোদন প্রয়োজন হবে এবং অস্থায়ী বহিষ্কারের ক্ষেত্রে সংগঠনের মোট সদস্যদের দুই-তৃতীয়াংশ (২/৩) এর অনুমোদন প্রয়োজন হবে। এখানে উল্লেখ থাকে, অস্থায়ী বহিষ্কারের সর্বনিম্ন মেয়াদকাল তিন (৩) মাস।", bold_prefix="ঙ.")
add_body_p(doc, " সদস্যপদ স্থায়ীভাবে বাতিল হওয়া অথবা সংগঠন থেকে স্থায়ী বহিষ্কার হওয়া কোনো সদস্য কার্যনির্বাহী কমিটির কোনো পদে অবস্থান করলেও তা স্বয়ংক্রিয়ভাবে বাতিল বলে গণ্য হবে এবং এরূপ ব্যক্তি আর সংগঠনের সদস্য নয় বিধায় কার্যনির্বাহী কমিটির কোনো পদে অধিষ্ঠিত হওয়ার সুযোগ থাকবে না।", bold_prefix="চ.", space_after=14)

add_section_title(doc, "ধারা-১১: সংশোধন ও হালনাগাদ (Amendment and Update)")
add_body_p(doc, " কার্যনির্বাহী কমিটি প্রয়োজনীয় পরিবর্তন, সংযোজন বা বিয়োজনসহ সংশোধিত খসড়া প্রণয়ন করবেন এবং উপদেষ্টামণ্ডলীর সদস্যদের কাছে উপস্থাপন করবেন।", bold_prefix="১.")
add_body_p(doc, " উপদেষ্টামণ্ডলীর সদস্যদের দুই-তৃতীয়াংশ (২/৩) এর অনুমোদনের ভিত্তিতে খসড়া গঠনতন্ত্র চূড়ান্ত করা হবে।", bold_prefix="২.")
add_body_p(doc, " উপদেষ্টামণ্ডলীর কোনো সদস্য সাত (৭) কার্যদিবসের মধ্যে অনাপত্তি না জানালে, তা স্বয়ংক্রিয়ভাবে অনুমোদিত হিসেবে গণ্য হবে।", bold_prefix="৩.")
add_body_p(doc, " খসড়া উপদেষ্টাদের কাছে উপস্থাপনের পর থেকে দশ (১০) কার্যদিবসের মধ্যে দুই-তৃতীয়াংশ (২/৩) এর অনুমোদনের ভিত্তিতে চূড়ান্ত করা হবে।", bold_prefix="৪.", space_after=14)

output_filename = "Munshiganj_District_Students_Association_Constitution.docx"
doc.save(output_filename)
print(f"Successfully generated {output_filename}")
